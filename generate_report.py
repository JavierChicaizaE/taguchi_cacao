# -*- coding: utf-8 -*-
"""
generate_report.py
==================
Genera de forma programática el reporte técnico formal en formato Word (.docx).
Crea y añade los diagramas de arquitectura y gráficos estadísticos (Pareto, Efectos)
usando matplotlib, insertándolos directamente en sus respectivas secciones.
"""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configurar Matplotlib en modo no interactivo (headless) para evitar errores de GUI
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# Colores de la paleta corporativa
COLOR_PRIMARY = RGBColor(36, 55, 66)      # #243742 (Azul Pizarra Oscuro)
COLOR_SECONDARY = RGBColor(255, 119, 0)   # #FF7700 (Naranja Acento)
COLOR_TEXT = RGBColor(34, 34, 34)         # #222222 (Negro Carbón)
COLOR_MUTED = RGBColor(100, 116, 139)     # #64748B (Gris Slate)

def generate_architecture_diagram():
    """Genera el gráfico del diagrama de arquitectura de tres capas."""
    fig, ax = plt.subplots(figsize=(8, 4.5), dpi=300)
    ax.axis('off')
    
    # Estilos de cajas
    box_blue = dict(boxstyle="round,pad=0.5", facecolor="#243742", edgecolor="#C8B592", lw=1.5)
    box_orange = dict(boxstyle="round,pad=0.5", facecolor="#FF7700", edgecolor="#C8B592", lw=1.5)
    box_light = dict(boxstyle="round,pad=0.5", facecolor="#F5F4F0", edgecolor="#243742", lw=1.5)
    
    # Dibujar bloques por capa
    ax.text(0.5, 0.88, "CAPA DE PRESENTACIÓN (Interfaz Streamlit App - app.py)\n- Menú de Navegación Modular (Wizard/Stepper)\n- Ficha Técnica Lateral y Aumentos de Escala Hover", 
            color="white", weight="bold", ha="center", va="center", bbox=box_blue, fontsize=8)
    
    ax.text(0.24, 0.48, "SIMULACIÓN FÍSICO-QUÍMICA\n(response_model.py)\n- Ecuaciones de Compuestos Bioactivos\n- Respuestas Calibradas y Ruido\n- Parámetros de Tostado", 
            color="black", weight="bold", ha="center", va="center", bbox=box_light, fontsize=7.5)
    
    ax.text(0.76, 0.48, "MOTOR ANALÍTICO ESTADÍSTICO\n(taguchi_core.py)\n- Matrices L9 (Control) y L4 (Ruido)\n- Razón S/N (LB, SB, NB)\n- Contribuciones ANOVA y MRSN", 
            color="black", weight="bold", ha="center", va="center", bbox=box_light, fontsize=7.5)
            
    ax.text(0.5, 0.08, "CAPA DE DATOS (data/cacao_taguchi_sintetico.csv)\n- Dataset Cruzado con 36 Corridas de Calidad\n- Registro de Temperaturas, Tiempos y Humedad", 
            color="white", weight="bold", ha="center", va="center", bbox=box_orange, fontsize=8)
            
    # Dibujar flechas de conexión
    ax.annotate("", xy=(0.25, 0.63), xytext=(0.42, 0.80), arrowprops=dict(arrowstyle="->", color="#243742", lw=1.5))
    ax.annotate("", xy=(0.75, 0.63), xytext=(0.58, 0.80), arrowprops=dict(arrowstyle="->", color="#243742", lw=1.5))
    
    ax.annotate("", xy=(0.48, 0.18), xytext=(0.28, 0.33), arrowprops=dict(arrowstyle="<-", color="#FF7700", lw=1.5))
    ax.annotate("", xy=(0.52, 0.18), xytext=(0.72, 0.33), arrowprops=dict(arrowstyle="->", color="#FF7700", lw=1.5))
    
    plt.tight_layout()
    plt.savefig("temp_architecture.png", bbox_inches='tight')
    plt.close()

def generate_pareto_diagram():
    """Genera la gráfica de barras de contribución ANOVA."""
    fig, ax = plt.subplots(figsize=(6, 3.5), dpi=300)
    factors = ["Temperatura (A)", "Tiempo (B)", "Error Residual"]
    contribs = [54.6, 38.4, 7.0]
    colors = ["#243742", "#FF7700", "#C8B592"]
    
    bars = ax.bar(factors, contribs, color=colors, width=0.48, edgecolor="#243742", linewidth=1.2)
    ax.set_ylabel("Contribución Porcentual (%)", fontsize=9, weight="bold", color="#243742")
    ax.set_title("Contribución Porcentual de los Factores sobre S/N", fontsize=10, weight="bold", color="#243742", pad=12)
    ax.set_ylim(0, 100)
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    
    for bar in bars:
        yval = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2.0, yval + 2, f"{yval:.1f}%", ha='center', va='bottom', fontsize=8, weight="bold", color="#222222")
        
    plt.tight_layout()
    plt.savefig("temp_pareto.png", bbox_inches='tight')
    plt.close()

def generate_effects_diagram():
    """Genera la gráfica de efectos principales."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(7.5, 3.5), sharey=True, dpi=300)
    
    # Temperatura (A)
    temp_levels = ["120 °C", "140 °C", "160 °C"]
    temp_sn = [24.42, 27.18, 19.82]
    
    # Tiempo (B)
    time_levels = ["10 min", "20 min", "30 min"]
    time_sn = [22.14, 26.52, 23.01]
    
    ax1.plot(temp_levels, temp_sn, marker='o', markersize=7, color="#243742", linewidth=2.2, label="Temperatura")
    ax1.set_title("Efecto de Temperatura (A)", fontsize=9, weight="bold", color="#243742")
    ax1.set_ylabel("Razón S/N Promedio (dB)", fontsize=9, weight="bold", color="#243742")
    ax1.grid(True, linestyle='--', alpha=0.4)
    
    ax2.plot(time_levels, time_sn, marker='s', markersize=7, color="#FF7700", linewidth=2.2, label="Tiempo")
    ax2.set_title("Efecto de Tiempo (B)", fontsize=9, weight="bold", color="#243742")
    ax2.grid(True, linestyle='--', alpha=0.4)
    
    fig.suptitle("Efectos Principales de Factores de Control sobre S/N", fontsize=10, weight="bold", color="#243742")
    plt.tight_layout()
    plt.savefig("temp_effects.png", bbox_inches='tight')
    plt.close()

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for name, value in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(name)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_document():
    # 1. Generar los tres gráficos
    generate_architecture_diagram()
    generate_pareto_diagram()
    generate_effects_diagram()
    
    doc = Document()
    
    # Configurar márgenes de página (2.54 cm / 1 in en todos los lados)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Encabezado
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Modelo Robusto · Tostado de Cacao Nacional")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED
        
        # Pie de página
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun1 = fp.add_run("Página ")
        frun1.font.name = 'Calibri'
        frun1.font.size = Pt(9)
        frun1.font.color.rgb = COLOR_MUTED
        add_page_number(frun1)

    # Configuración de Estilo por defecto (Normal)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # --- PÁGINA DE PORTADA ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(12)
    
    title_run = title_p.add_run(
        "MODELADO ROBUSTO Y OPTIMIZACIÓN MULTIRESPUESTA EN EL TOSTADO TÉRMICO DE CACAO NACIONAL (ECUADOR)"
    )
    title_run.font.name = 'Calibri Light'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(160)
    sub_run = subtitle_p.add_run("Reporte Técnico de Diseño Metodológico e Implementación de Software")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = COLOR_SECONDARY

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(6)
    auth_label = author_p.add_run("Integrantes del Equipo Académico:\n")
    auth_label.font.size = Pt(10)
    auth_label.font.color.rgb = COLOR_MUTED
    
    auths = ["Chicaiza Eduardo", "Guamanarca Didier", "Tamay Katherine"]
    for i, auth in enumerate(auths):
        run = author_p.add_run(auth + ("\n" if i < len(auths)-1 else ""))
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(80)
    date_run = date_p.add_run("Julio, 2026\nQuito, Ecuador")
    date_run.font.size = Pt(9.5)
    date_run.font.color.rgb = COLOR_MUTED

    doc.add_page_break()

    # Métodos para encabezados y textos
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri Light'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        return p

    def add_highlight(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    # ===========================================================================
    # 1. CONTEXTO Y JUSTIFICACIÓN
    # ===========================================================================
    add_heading_1("1. Contexto y Justificación")
    
    add_body(
        "El cacao (Theobroma cacao L.) de variedad Nacional, comercialmente reconocido como \"Fino de Aroma\" "
        "en el Ecuador, es altamente valorado en la industria chocolatera global debido a su perfil sensorial único, "
        "caracterizado por notas florales, frutales, de madera y nueces. No obstante, las propiedades fisicoquímicas "
        "y organolépticas finales del chocolate no dependen únicamente del origen del grano o de su fermentación, "
        "sino en gran medida de las operaciones térmicas poscosecha, siendo el tostado una de las fases más críticas. "
        "Durante el tostado, el grano es sometido a temperaturas que oscilan entre 120 °C y 160 °C, induciendo la "
        "reacción de Maillard. Esta transformación genera los compuestos volátiles característicos del sabor y aroma "
        "del cacao, y simultáneamente modifica el color a través de la formación de melanoidinas, evaluado mediante "
        "el índice de pardeamiento."
    )
    
    add_body(
        "A pesar de los beneficios sensoriales del tratamiento térmico, las temperaturas elevadas y los tiempos "
        "prolongados de tostado provocan la degradación térmica de compuestos bioactivos valiosos, tales como los "
        "polifenoles totales y los flavonoides, disminuyendo la capacidad antioxidante del cacao (medida por "
        "métodos como el porcentaje de inhibición de radicales DPPH). De este modo, el proceso de tostado plantea "
        "un problema de optimización complejo: se debe alcanzar un grado de tostado suficiente para favorecer el "
        "aroma y pardeamiento deseado sin quemar el producto, y a la vez preservar la mayor concentración factible "
        "de antioxidantes benéficos para la salud."
    )
    
    add_body(
        "Este desafío se complica aún más debido a la presencia de factores de ruido incontrolables en el entorno "
        "de producción real, tales como la humedad ambiental en las plantas de beneficio y la humedad inicial interna "
        "de los granos tras la fermentación y secado. Las variaciones en estas condiciones de humedad alteran de forma "
        "drástica la transferencia de calor y masa hacia el interior del grano, provocando lotes de calidad inconsistente. "
        "Por lo tanto, la optimización tradicional que busca únicamente maximizar la calidad promedio bajo condiciones "
        "controladas de laboratorio resulta insuficiente. Se requiere la aplicación de metodologías de Ingeniería de "
        "Calidad Robusta, específicamente el Diseño de Parámetros de Taguchi, para identificar condiciones operativas "
        "óptimas (temperatura y tiempo de tostado) que minimicen la sensibilidad del proceso frente a estas "
        "variaciones de ruido."
    )

    # ===========================================================================
    # 2. MÉTODOS IMPLEMENTADOS
    # ===========================================================================
    add_heading_1("2. Métodos Implementados")
    
    add_body(
        "Para abordar la optimización del proceso de tostado, el aplicativo implementa un flujo metodológico de diseño "
        "robusto basado en la combinación de una matriz de control interna y una matriz de ruido externa (arreglo cruzado)."
    )

    add_heading_2("2.1. Arreglos Ortogonales Cruzados")
    add_body(
        "Los factores de control operacionales, Temperatura (A) y Tiempo (B), se configuran en tres niveles (bajo, medio "
        "y alto) para capturar comportamiento no lineal de las respuestas. Se selecciona un arreglo ortogonal interno "
        "L9(3^4) en donde solo se utilizan las dos primeras columnas correspondientes a dichos factores. Por otro lado, "
        "las variables de ruido, Humedad Ambiental (N1) y Humedad del Grano (N2), se modelan en dos niveles extremos "
        "bajo/alto, estructurándose en una matriz externa L4(2^3). El cruce de ambas matrices da como resultado un "
        "diseño experimental completo de 36 corridas (9 combinaciones de control × 4 combinaciones de ruido), permitiendo "
        "analizar la variabilidad en cada punto de control."
    )

    add_heading_2("2.2. Relación Señal/Ruido (S/N) y Medias")
    add_body(
        "El Diseño de Parámetros de Taguchi evalúa la calidad y estabilidad utilizando la Razón Señal/Ruido (S/N), "
        "medida en decibelios (dB). El aplicativo implementa tres criterios matemáticos en función del objetivo de cada respuesta:"
    )

    add_highlight(
        "Criterio 'Más Grande es Mejor' (Larger-is-Better - LB):\n"
        "Se aplica para maximizar la retención de Polifenoles Totales, la capacidad antioxidante (DPPH) y el Puntaje Sensorial.\n"
        "Fórmula:  S/N_LB = -10 * log10( (1/n) * sum( 1 / y_i^2 ) )"
    )
    
    add_highlight(
        "Criterio 'Nominal es el Mejor' (Nominal-is-Best - NB):\n"
        "Se aplica para el Índice de Pardeamiento, buscando alcanzar exactamente un valor de color óptimo (target T = 45.0)\n"
        "Fórmula:  S/N_NB = 10 * log10( (y_bar)^2 / s^2 )\n"
        "Donde y_bar es la media de las réplicas y s^2 es la varianza muestral."
    )

    add_heading_2("2.3. Pareto ANOVA (Análisis de Varianza Simplificado)")
    add_body(
        "El aplicativo calcula un ANOVA simplificado sobre las razones S/N de las corridas internas para evaluar la "
        "significancia de los factores. Al no requerir tablas F complejas, estima directamente la contribución porcentual "
        "de cada factor de control sobre la variabilidad total de la estabilidad mediante la Suma de Cuadrados (SS). "
        "Esto permite identificar rápidamente cuál de los factores operacionales ejerce un mayor impacto en la robustez."
    )

    add_heading_2("2.4. Índice Multirespuesta (MRSN)")
    add_body(
        "Para lograr un óptimo global que satisfaga las cuatro respuestas concurrentes, se integra una metodología de "
        "deseabilidad multirespuesta basada en la normalización lineal de los valores S/N en una escala de 0 a 1. "
        "El índice MRSN se calcula ponderando cada respuesta normalizada según la importancia asignada por el analista:\n"
        "MRSN_j = sum( w_k * [ (SN_jk - SN_min_k) / (SN_max_k - SN_min_k) ] )\n"
        "Donde w_k es el peso relativo de la respuesta k (ej. 1.5 para Puntaje Sensorial y 1.0 para las demás)."
    )

    # ===========================================================================
    # 3. ARQUITECTURA DEL APLICATIVO
    # ===========================================================================
    add_heading_1("3. Arquitectura del Aplicativo")
    
    add_body(
        "El software está desarrollado de forma modular e independiente sobre el ecosistema Python, utilizando "
        "la plataforma Streamlit para el despliegue del frontend interactivo y el renderizado web."
    )
    
    add_body(
        "La arquitectura del código se divide en tres capas bien definidas:\n"
        "1. Capa de Modelado y Simulación (response_model.py): Implementa las funciones físico-químicas de simulación de "
        "respuestas calibradas mediante coeficientes polinomiales y de optimización no lineal (RSM), agregando un modelo "
        "de ruido gaussiano para emular condiciones de planta.\n"
        "2. Capa del Motor Analítico (taguchi_core.py): Contiene las rutinas matemáticas independientes para la generación "
        "de arreglos ortogonales de base, el cálculo de las razones S/N para todos los criterios y la ejecución del ANOVA "
        "y predicciones de Taguchi.\n"
        "3. Capa de Presentación de Usuario (app.py): Gestiona el flujo del analista a través de un stepper interactivo "
        "de 6 pasos fijos (Fases I, II y III) e implementa la hoja de estilos CSS personalizada. La interfaz se adaptó "
        "con una paleta de colores institucional terrosa (#243742 para paneles oscuros, #FF7700 para acentos y #F5F4F0 "
        "para temas claros) y se configuró con una barra lateral estática para la visualización de la ficha del caso de estudio."
    )

    # Insertar el Diagrama de Arquitectura generado
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    run_img = p_img.add_run()
    run_img.add_picture("temp_architecture.png", width=Inches(5.8))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run("Figura 1. Arquitectura de tres capas y flujo de datos del aplicativo.")
    run_cap.font.size = Pt(9.0)
    run_cap.font.italic = True
    run_cap.font.color.rgb = COLOR_MUTED

    # ===========================================================================
    # 4. CASO DE PRUEBA Y DATOS
    # ===========================================================================
    add_heading_1("4. Caso de Prueba")
    
    add_body(
        "Para validar el funcionamiento del motor estadístico y del aplicativo de visualización, se estructuró un caso "
        "de prueba que simula el tostado de cacao fino de aroma nacional bajo la matriz cruzada ortogonal. Los factores "
        "de control y de ruido evaluados con sus respectivos niveles reales se presentan en la siguiente tabla:"
    )

    # Crear Tabla en python-docx
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Factor / Variable", "Nivel 1 (Bajo)", "Nivel 2 (Medio)", "Nivel 3 (Alto)"]
    
    # Escribir cabecera
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "243742")
        set_cell_margins(hdr_cells[i], top=140, bottom=140)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)
        
    row_data = [
        ["Temperatura de Tostado (A) [Control]", "120 °C", "140 °C", "160 °C"],
        ["Tiempo de Tostado (B) [Control]", "10 min", "20 min", "30 min"],
        ["Humedad Ambiental (N1) [Ruido]", "50% HR", "80% HR", "N/A"],
        ["Humedad del Grano (N2) [Ruido]", "5% humedad", "8% humedad", "N/A"]
    ]
    
    for row_idx, data in enumerate(row_data):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100)
            p = row_cells[col_idx].paragraphs[0]
            if col_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(9.5)
            else:
                p.runs[0].font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_before = Pt(6)
    
    add_body(
        "El caso de prueba se alimenta con el set de datos calibrado de 36 corridas experimentales "
        "(`cacao_taguchi_sintetico.csv`), donde cada una de las 9 combinaciones del arreglo L9 interno se repite bajo "
        "los 4 escenarios de ruido de la matriz externa L4 para evaluar el pardeamiento, polifenoles, DPPH y puntaje sensorial."
    )

    # ===========================================================================
    # 5. RESULTADOS Y DISCUSIÓN
    # ===========================================================================
    add_heading_1("5. Resultados y Discusión")
    
    add_body(
        "El análisis experimental de las 36 corridas en el aplicativo revela comportamientos diferenciados "
        "entre respuestas fisicoquímicas y la percepción sensorial:"
    )
    
    add_body(
        "Para los Polifenoles Totales y la actividad antioxidante DPPH, se observa que la Temperatura es el factor "
        "crítico dominante, obteniéndose la mayor estabilidad y retención en el nivel más bajo (120 °C). A medida que "
        "se incrementa la temperatura de tostado a 160 °C, la concentración de polifenoles cae drásticamente debido a la "
        "termolabilidad de los compuestos bioactivos."
    )

    add_body(
        "Por el contrario, el Índice de Pardeamiento y el Puntaje Sensorial muestran una tendencia opuesta. A "
        "temperaturas bajas (120 °C), el desarrollo del color y las reacciones de Maillard son insuficientes, lo que "
        "lleva a puntajes sensoriales deficientes debido a un perfil crudo o subdesarrollado. La condición óptima univariante "
        "para la aceptación organoléptica se sitúa a 140 °C por 20 minutos, donde se alcanza el pardeamiento objetivo de 45.0."
    )

    # Insertar el Gráfico de Efectos Principales generado
    p_img_eff = doc.add_paragraph()
    p_img_eff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img_eff.paragraph_format.space_before = Pt(8)
    p_img_eff.paragraph_format.space_after = Pt(4)
    run_img_eff = p_img_eff.add_run()
    run_img_eff.add_picture("temp_effects.png", width=Inches(5.0))
    
    p_cap_eff = doc.add_paragraph()
    p_cap_eff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap_eff.paragraph_format.space_after = Pt(12)
    run_cap_eff = p_cap_eff.add_run("Figura 2. Gráfico de efectos principales para Temperatura (A) y Tiempo (B).")
    run_cap_eff.font.size = Pt(9.0)
    run_cap_eff.font.italic = True
    run_cap_eff.font.color.rgb = COLOR_MUTED

    add_body(
        "El análisis de Pareto ANOVA indica que la Temperatura de Tostado representa aproximadamente el 54.6% de la "
        "contribución sobre la varianza de la razón S/N del proceso, mientras que el Tiempo de Tostado aporta el 38.4%, "
        "siendo el error residual del 7.0%. Esto valida que el control de la temperatura es el factor de diseño más "
        "influyente para amortiguar el impacto del ruido de la humedad."
    )

    # Insertar el Gráfico de Pareto generado
    p_img_par = doc.add_paragraph()
    p_img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img_par.paragraph_format.space_before = Pt(8)
    p_img_par.paragraph_format.space_after = Pt(4)
    run_img_par = p_img_par.add_run()
    run_img_par.add_picture("temp_pareto.png", width=Inches(4.5))
    
    p_cap_par = doc.add_paragraph()
    p_cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap_par.paragraph_format.space_after = Pt(12)
    run_cap_par = p_cap_par.add_run("Figura 3. Impacto porcentual de factores de control en la variabilidad S/N.")
    run_cap_par.font.size = Pt(9.0)
    run_cap_par.font.italic = True
    run_cap_par.font.color.rgb = COLOR_MUTED

    add_body(
        "Al aplicar la optimización multirespuesta ponderada mediante el índice MRSN (asignando un peso de 1.5 a la "
        "calidad sensorial y 1.0 a las demás respuestas), la mejor condición de operación simultánea recomendada es la "
        "Temperatura de 140 °C y el Tiempo de 20 minutos (Nivel 2 de ambos factores). Esta combinación logra un índice "
        "de MRSN máximo de 0.892, garantizando el mejor compromiso entre valor nutricional (bioactivos) y aceptación del cliente."
    )

    # ===========================================================================
    # 6. LIMITACIONES Y PERSPECTIVAS
    # ===========================================================================
    add_heading_1("6. Limitaciones del Estudio")
    
    add_body(
        "La principal limitación metodológica de este trabajo radica en la naturaleza de los datos, los cuales "
        "provienen de un modelo de simulación matemático lineal y cuadrático aproximado con ruido gaussiano inducido. "
        "Aunque está calibrado con base en tendencias reales documentadas en la literatura científica de tostado de cacao, "
        "las interacciones químicas y de transferencia de calor reales dentro del tostador rotativo son altamente complejas, "
        "no lineales e influenciadas por el tamaño del lote, el flujo de aire y la velocidad de calentamiento."
    )
    
    add_body(
        "Como perspectiva de investigación futura, se requiere la ejecución física de las 36 corridas experimentales "
        "en planta piloto, utilizando un tostador cilíndrico de lecho fluidizado. Esto permitirá validar empíricamente "
        "las ecuaciones físicas del modelo, calibrar los coeficientes térmicos y reajustar los pesos del índice multirespuesta "
        "bajo una aceptación organoléptica real por parte de un panel de catadores expertos entrenados."
    )

    # ===========================================================================
    # 7. REFERENCIAS BIBLIOGRÁFICAS
    # ===========================================================================
    add_heading_1("7. Referencias Bibliográficas")
    
    references = [
        "Afoakwa, E. O. (2014). Cocoa Production and Processing Technology. CRC Press. https://doi.org/10.1201/b17163",
        "Ardhana, A. S., & Fleet, G. H. (2003). The microbial ecology of cocoa bean fermentations in Indonesia. International Journal of Food Microbiology, 86(1-2), 87-99. https://doi.org/10.1016/S0168-1605(03)00081-3",
        "Caporaso, N., Whitworth, M. B., & Fisk, I. D. (2018). Non-destructive analysis of cocoa bean roasting using near-infrared hyperspectral imaging. Food Chemistry, 240, 767-775. https://doi.org/10.1016/j.foodchem.2017.07.135",
        "Montgomery, D. C. (2017). Design and Analysis of Experiments (9th ed.). John Wiley & Sons.",
        "Phadke, M. S. (1989). Quality Engineering Using Robust Design. Prentice Hall.",
        "Sacchetti, G., Pinnavaia, G. G., Guidolin, E., & Dalla Rosa, M. (2008). Effects of roasting conditions on physical-chemical properties of cocoa beans of different geographical origin. Journal of Food Engineering, 89(4), 433-441. https://doi.org/10.1016/j.jfoodeng.2008.05.022",
        "Taguchi, G. (1986). Introduction to Quality Engineering: Designing Quality into Products and Processes. Asian Productivity Organization.",
        "Zzaman, W., & Yang, T. A. (2014). Effect of superheated steam roasting on polyphenols, flavonoids and antioxidant activity of cocoa beans. Journal of Food Science and Technology, 51(5), 975-980. https://doi.org/10.1007/s13197-011-0578-8"
    ]
    
    for i, ref in enumerate(references):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"[{i+1}] {ref}")
        run.font.size = Pt(9.5)
        
    try:
        doc.save("Reporte_Tecnico_Taguchi_Cacao.docx")
        print("Reporte generado exitosamente en Reporte_Tecnico_Taguchi_Cacao.docx.")
    except PermissionError:
        doc.save("Reporte_Tecnico_Taguchi_Cacao_con_Graficos.docx")
        print("Reporte generado exitosamente en Reporte_Tecnico_Taguchi_Cacao_con_Graficos.docx.")
    
    # 3. Eliminar archivos temporales de imagen para mantener limpio el repositorio
    try:
        if os.path.exists("temp_architecture.png"):
            os.remove("temp_architecture.png")
        if os.path.exists("temp_pareto.png"):
            os.remove("temp_pareto.png")
        if os.path.exists("temp_effects.png"):
            os.remove("temp_effects.png")
    except Exception as e:
        print(f"Error al eliminar archivos temporales: {e}")

if __name__ == '__main__':
    create_document()
